#!/usr/bin/env python3
"""
njcivil.py - Terminal utility to:
  1) Extract searchable text from NJ eCourts/JEDS filings (PDF/DOCX) into .txt
  2) Build a SQLite FTS index for fast keyword search
  3) Generate draft pleadings (complaints, motions, certifications) from YAML facts

Designed for pro se NJ Superior Court (Law Division - Civil Part) workflows.
"""

from __future__ import annotations

import json
import os
import re
import sqlite3
import sys
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable, Optional, Tuple, List

import typer
from jinja2 import Environment, FileSystemLoader, StrictUndefined

# Optional imports; tool still works without them (with reduced capability).
try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None  # type: ignore

try:
    from docx import Document
except Exception:
    Document = None  # type: ignore

try:
    import yaml
except Exception:
    yaml = None  # type: ignore

try:
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:
    Pt = None  # type: ignore
    WD_ALIGN_PARAGRAPH = None  # type: ignore

app = typer.Typer(add_completion=False, help="NJ Civil terminal helper: extract/search filings; generate drafts.")


def _safe_read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _pdf_to_text_pypdf(pdf_path: Path) -> str:
    if PdfReader is None:
        raise RuntimeError("pypdf is not available. Install with: pip install pypdf")
    reader = PdfReader(str(pdf_path))
    parts: List[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            parts.append("")
    return "\n".join(parts)


def _docx_to_text_python_docx(docx_path: Path) -> str:
    if Document is None:
        raise RuntimeError("python-docx is not available. Install with: pip install python-docx")
    doc = Document(str(docx_path))
    parts: List[str] = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    # tables (common in captions)
    for t in doc.tables:
        for row in t.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


def _normalize(text: str) -> str:
    # Remove common footer/header noise
    text = re.sub(r"Page\s+\d+\s+of\s+\d+", "", text, flags=re.I)
    text = re.sub(r"\s+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


@app.command()
def extract(
    in_dir: Path = typer.Argument(..., exists=True, file_okay=False, dir_okay=True),
    out_dir: Path = typer.Argument(..., file_okay=False, dir_okay=True),
    glob: str = typer.Option("*.pdf;*.docx", help="Semicolon-separated glob patterns."),
    overwrite: bool = typer.Option(False, help="Overwrite existing .txt outputs."),
):
    """
    Extract text from PDF/DOCX into out_dir as .txt and write index.jsonl metadata.
    """
    out_dir.mkdir(parents=True, exist_ok=True)
    patterns = [g.strip() for g in glob.split(";") if g.strip()]
    paths: List[Path] = []
    for pat in patterns:
        paths.extend(sorted(in_dir.rglob(pat)))
    if not paths:
        raise typer.Exit(code=2)

    index_path = out_dir / "index.jsonl"
    mode = "w" if overwrite or not index_path.exists() else "a"
    with index_path.open(mode, encoding="utf-8") as idx:
        for p in paths:
            out_txt = out_dir / (p.name + ".txt")
            if out_txt.exists() and not overwrite:
                continue
            try:
                if p.suffix.lower() == ".pdf":
                    txt = _pdf_to_text_pypdf(p)
                elif p.suffix.lower() == ".docx":
                    txt = _docx_to_text_python_docx(p)
                else:
                    continue
                txt = _normalize(txt)
                out_txt.write_text(txt, encoding="utf-8")
                meta = {
                    "source": str(p),
                    "out_txt": str(out_txt),
                    "bytes": p.stat().st_size,
                    "chars": len(txt),
                    "updated_at": datetime.now().isoformat(timespec="seconds"),
                }
                idx.write(json.dumps(meta, ensure_ascii=False) + "\n")
                typer.echo(f"OK: {p.name}")
            except Exception as e:
                typer.echo(f"FAIL: {p.name} :: {e}", err=True)


def _iter_txt_files(txt_dir: Path) -> Iterable[Path]:
    for p in sorted(txt_dir.rglob("*.txt")):
        if p.name.endswith(".jsonl"):
            continue
        yield p


@app.command()
def grep(
    txt_dir: Path = typer.Argument(..., exists=True, file_okay=False, dir_okay=True),
    pattern: str = typer.Argument(..., help="Regex or plain text (use --literal to escape)."),
    literal: bool = typer.Option(False, help="Treat pattern as literal text."),
    context: int = typer.Option(80, help="Context characters before/after match."),
    max_hits: int = typer.Option(50, help="Maximum matches to print."),
):
    """
    Search extracted text files and print matches with context.
    """
    rx = re.compile(re.escape(pattern) if literal else pattern, flags=re.I)
    hits = 0
    for p in _iter_txt_files(txt_dir):
        t = _safe_read_text(p)
        for m in rx.finditer(t):
            hits += 1
            if hits > max_hits:
                raise typer.Exit(code=0)
            s = max(0, m.start() - context)
            e = min(len(t), m.end() + context)
            snippet = t[s:e].replace("\n", " ")
            typer.echo(f"\n[{p.name}] ...{snippet}...")
    if hits == 0:
        raise typer.Exit(code=1)


@app.command()
def index(
    txt_dir: Path = typer.Argument(..., exists=True, file_okay=False, dir_okay=True),
    db_path: Path = typer.Argument(..., help="SQLite db file to create/overwrite."),
):
    """
    Build a SQLite FTS5 full-text index from extracted .txt files.
    """
    if db_path.exists():
        db_path.unlink()
    con = sqlite3.connect(str(db_path))
    cur = con.cursor()
    cur.execute("PRAGMA journal_mode=WAL;")
    cur.execute("CREATE VIRTUAL TABLE docs USING fts5(path, title, body);")
    for p in _iter_txt_files(txt_dir):
        body = _safe_read_text(p)
        cur.execute("INSERT INTO docs(path, title, body) VALUES(?, ?, ?);", (str(p), p.name, body))
    con.commit()
    con.close()
    typer.echo(f"Indexed to {db_path}")


@app.command()
def query(
    db_path: Path = typer.Argument(..., exists=True, dir_okay=False),
    q: str = typer.Argument(..., help='FTS query, e.g., "tow NEAR/5 notice" or "Ruiz".'),
    limit: int = typer.Option(20, help="Max results."),
):
    """
    Query the SQLite FTS index and print ranked hits.
    """
    con = sqlite3.connect(str(db_path))
    cur = con.cursor()
    cur.execute("SELECT title, snippet(docs, 2, '[', ']', ' … ', 12) as snip FROM docs WHERE docs MATCH ? LIMIT ?;", (q, limit))
    rows = cur.fetchall()
    con.close()
    if not rows:
        raise typer.Exit(code=1)
    for title, snip in rows:
        typer.echo(f"\n{title}\n{snip}")


def _load_yaml(path: Path) -> dict:
    if yaml is None:
        raise RuntimeError("PyYAML not available. Install with: pip install pyyaml")
    return yaml.safe_load(path.read_text(encoding="utf-8"))


def _render_template(templates_dir: Path, template_name: str, data: dict) -> str:
    env = Environment(
        loader=FileSystemLoader(str(templates_dir)),
        undefined=StrictUndefined,
        autoescape=False,
        trim_blocks=True,
        lstrip_blocks=True,
    )
    tpl = env.get_template(template_name)
    return tpl.render(**data).strip() + "\n"


def _docx_from_text(out_path: Path, title: str, text: str):
    if Document is None or Pt is None:
        raise RuntimeError("python-docx not available. Install with: pip install python-docx")
    doc = Document()
    # Basic font (Word defaults vary; set to Times New Roman 12 for court readability)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Title (optional)
    if title:
        p = doc.add_paragraph(title)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for line in text.splitlines():
        line = line.rstrip()
        if not line:
            doc.add_paragraph("")  # blank line
            continue
        p = doc.add_paragraph(line)
        # Bold headings heuristics
        if re.fullmatch(r"(SUPERIOR COURT OF NEW JERSEY|LAW DIVISION.*|[IVX]+\.\s+.*|PRELIMINARY STATEMENT|PARTIES|JURISDICTION AND VENUE|FACTUAL ALLEGATIONS|PRAYER FOR RELIEF|JURY DEMAND|VERIFICATION.*|CERTIFICATION.*|COUNT\s+[IVX]+.*)", line.strip(), flags=re.I):
            for run in p.runs:
                run.bold = True
    doc.save(str(out_path))


@app.command()
def make(
    kind: str = typer.Argument(..., help="Template kind: verified_complaint | motion_interim | certification"),
    facts: Path = typer.Option(..., exists=True, dir_okay=False, help="YAML file with case facts/party info."),
    out: Path = typer.Option(..., help="Output path (.txt or .docx)."),
    templates_dir: Path = typer.Option(None, help="Templates directory (defaults to ./templates next to script)."),
):
    """
    Generate a draft pleading from templates using YAML facts.

    Examples:
      python njcivil.py make verified_complaint --facts facts.yml --out complaint.docx
      python njcivil.py make certification --facts facts.yml --out cert.txt
    """
    templates_dir = templates_dir or (Path(__file__).resolve().parent / "templates")
    data = _load_yaml(facts)
    template_map = {
        "verified_complaint": "verified_complaint.j2",
        "motion_interim": "motion_interim_relief.j2",
        "certification": "certification_support.j2",
    }
    if kind not in template_map:
        raise typer.BadParameter(f"Unknown kind: {kind}. Options: {', '.join(template_map)}")

    rendered = _render_template(templates_dir, template_map[kind], data)
    out.parent.mkdir(parents=True, exist_ok=True)

    if out.suffix.lower() == ".docx":
        _docx_from_text(out, title="", text=rendered)
    else:
        out.write_text(rendered, encoding="utf-8")
    typer.echo(f"Wrote {out}")


@app.command()
def merge_pdfs(
    out_pdf: Path = typer.Argument(..., help="Output PDF path"),
    pdfs: List[Path] = typer.Argument(..., help="Input PDFs to merge in order"),
):
    """
    Merge PDFs into one (useful for exhibits).
    """
    if PdfReader is None:
        raise RuntimeError("pypdf not available. Install with: pip install pypdf")
    from pypdf import PdfWriter
    writer = PdfWriter()
    for p in pdfs:
        r = PdfReader(str(p))
        for page in r.pages:
            writer.add_page(page)
    out_pdf.parent.mkdir(parents=True, exist_ok=True)
    with out_pdf.open("wb") as f:
        writer.write(f)
    typer.echo(f"Merged to {out_pdf}")


if __name__ == "__main__":
    app()

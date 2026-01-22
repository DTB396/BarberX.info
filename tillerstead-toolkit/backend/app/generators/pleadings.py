"""
BarberX Legal Case Management Pro Suite
NJ Civil Pleadings Generator
Integrates njcivil_terminal_tool for document generation
"""
import json
import os
import re
import sqlite3
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
from enum import Enum

# Optional imports
try:
    from pypdf import PdfReader, PdfWriter
except ImportError:
    PdfReader = None
    PdfWriter = None

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None
    Pt = None
    WD_ALIGN_PARAGRAPH = None

try:
    import yaml
except ImportError:
    yaml = None

try:
    from jinja2 import Environment, FileSystemLoader, StrictUndefined, BaseLoader
except ImportError:
    Environment = None


class PleadingType(Enum):
    """Types of NJ Civil pleadings"""
    VERIFIED_COMPLAINT = "verified_complaint"
    MOTION_INTERIM_RELIEF = "motion_interim_relief"
    CERTIFICATION_SUPPORT = "certification_support"
    ANSWER = "answer"
    MOTION_TO_DISMISS = "motion_to_dismiss"
    MOTION_TO_COMPEL = "motion_to_compel"
    MOTION_FOR_SUMMARY_JUDGMENT = "motion_for_summary_judgment"
    MOTION_TO_SUPPRESS = "motion_to_suppress"
    BRIEF_IN_SUPPORT = "brief_in_support"
    BRIEF_IN_OPPOSITION = "brief_in_opposition"
    DISCOVERY_INTERROGATORIES = "discovery_interrogatories"
    DISCOVERY_RFP = "discovery_rfp"
    DISCOVERY_RFA = "discovery_rfa"
    SUBPOENA = "subpoena"
    NOTICE_OF_MOTION = "notice_of_motion"
    ORDER_TO_SHOW_CAUSE = "order_to_show_cause"


@dataclass
class PleadingResult:
    """Result of pleading generation"""
    success: bool
    pleading_type: PleadingType
    output_path: Optional[str]
    content: Optional[str]
    error: Optional[str]
    metadata: Dict[str, Any]


@dataclass  
class SearchResult:
    """Result from full-text search"""
    document: str
    snippet: str
    score: float
    context: str


class NJCivilPleadingsGenerator:
    """
    NJ Civil Pleadings Generator
    
    Generates court documents from templates using case facts.
    Supports YAML fact files and direct dictionary input.
    """
    
    # Template mapping
    TEMPLATE_MAP = {
        PleadingType.VERIFIED_COMPLAINT: "verified_complaint.j2",
        PleadingType.MOTION_INTERIM_RELIEF: "motion_interim_relief.j2",
        PleadingType.CERTIFICATION_SUPPORT: "certification_support.j2",
    }
    
    # Heading patterns for bold formatting in DOCX
    HEADING_PATTERNS = [
        r"SUPERIOR COURT OF NEW JERSEY",
        r"LAW DIVISION.*",
        r"[IVX]+\.\s+.*",
        r"PRELIMINARY STATEMENT",
        r"PARTIES",
        r"JURISDICTION AND VENUE",
        r"FACTUAL ALLEGATIONS",
        r"PRAYER FOR RELIEF",
        r"JURY DEMAND",
        r"VERIFICATION.*",
        r"CERTIFICATION.*",
        r"COUNT\s+[IVX]+.*",
        r"CLAIMS FOR RELIEF",
        r"LEGAL ARGUMENT",
        r"CONCLUSION",
        r"WHEREFORE",
    ]
    
    def __init__(self, templates_dir: Optional[Path] = None):
        """
        Initialize the generator.
        
        Args:
            templates_dir: Path to templates directory
        """
        if templates_dir is None:
            # Default to njcivil_terminal_tool/templates
            templates_dir = Path(__file__).parent.parent.parent / "njcivil_terminal_tool" / "templates"
        
        self.templates_dir = Path(templates_dir)
        self._heading_regex = re.compile(
            "|".join(f"({p})" for p in self.HEADING_PATTERNS),
            flags=re.IGNORECASE
        )
    
    def load_facts_yaml(self, yaml_path: Path) -> Dict[str, Any]:
        """Load facts from YAML file"""
        if yaml is None:
            raise RuntimeError("PyYAML not available. Install with: pip install pyyaml")
        
        return yaml.safe_load(yaml_path.read_text(encoding="utf-8"))
    
    def generate_pleading(
        self,
        pleading_type: PleadingType,
        facts: Dict[str, Any],
        output_path: Optional[Path] = None,
        output_format: str = "docx"
    ) -> PleadingResult:
        """
        Generate a pleading document.
        
        Args:
            pleading_type: Type of pleading to generate
            facts: Dictionary of case facts
            output_path: Optional output path
            output_format: Output format (docx, txt, html)
            
        Returns:
            PleadingResult with generated content
        """
        if Environment is None:
            return PleadingResult(
                success=False,
                pleading_type=pleading_type,
                output_path=None,
                content=None,
                error="Jinja2 not available. Install with: pip install jinja2",
                metadata={}
            )
        
        template_name = self.TEMPLATE_MAP.get(pleading_type)
        if not template_name:
            return PleadingResult(
                success=False,
                pleading_type=pleading_type,
                output_path=None,
                content=None,
                error=f"No template available for {pleading_type.value}",
                metadata={}
            )
        
        try:
            # Render template
            content = self._render_template(template_name, facts)
            
            # Save to file if output path provided
            if output_path:
                output_path = Path(output_path)
                output_path.parent.mkdir(parents=True, exist_ok=True)
                
                if output_format == "docx":
                    self._save_docx(output_path, content)
                elif output_format == "txt":
                    output_path.write_text(content, encoding="utf-8")
                elif output_format == "html":
                    html = self._text_to_html(content)
                    output_path.write_text(html, encoding="utf-8")
            
            return PleadingResult(
                success=True,
                pleading_type=pleading_type,
                output_path=str(output_path) if output_path else None,
                content=content,
                error=None,
                metadata={
                    "generated_at": datetime.now().isoformat(),
                    "template": template_name,
                    "format": output_format,
                    "word_count": len(content.split()),
                    "char_count": len(content)
                }
            )
            
        except Exception as e:
            return PleadingResult(
                success=False,
                pleading_type=pleading_type,
                output_path=None,
                content=None,
                error=str(e),
                metadata={}
            )
    
    def _render_template(self, template_name: str, data: Dict) -> str:
        """Render a Jinja2 template"""
        env = Environment(
            loader=FileSystemLoader(str(self.templates_dir)),
            undefined=StrictUndefined,
            autoescape=False,
            trim_blocks=True,
            lstrip_blocks=True,
        )
        tpl = env.get_template(template_name)
        return tpl.render(**data).strip() + "\n"
    
    def _save_docx(self, output_path: Path, text: str, title: str = ""):
        """Save content as DOCX with court formatting"""
        if Document is None or Pt is None:
            raise RuntimeError("python-docx not available")
        
        doc = Document()
        
        # Set default style
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        
        # Add title if provided
        if title:
            p = doc.add_paragraph(title)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Process lines
        for line in text.splitlines():
            line = line.rstrip()
            if not line:
                doc.add_paragraph("")
                continue
            
            p = doc.add_paragraph(line)
            
            # Bold headings
            if self._heading_regex.fullmatch(line.strip()):
                for run in p.runs:
                    run.bold = True
        
        doc.save(str(output_path))
    
    def _text_to_html(self, text: str) -> str:
        """Convert text to basic HTML"""
        lines = []
        lines.append("<!DOCTYPE html>")
        lines.append("<html><head>")
        lines.append("<style>")
        lines.append("body { font-family: 'Times New Roman', serif; font-size: 12pt; max-width: 8.5in; margin: 1in auto; }")
        lines.append("h1, h2, h3 { font-weight: bold; }")
        lines.append(".count { font-weight: bold; margin-top: 1em; }")
        lines.append("</style>")
        lines.append("</head><body>")
        
        for line in text.splitlines():
            line = line.strip()
            if not line:
                lines.append("<p>&nbsp;</p>")
            elif self._heading_regex.fullmatch(line):
                lines.append(f"<p><strong>{line}</strong></p>")
            else:
                lines.append(f"<p>{line}</p>")
        
        lines.append("</body></html>")
        return "\n".join(lines)
    
    def get_available_templates(self) -> List[Dict[str, Any]]:
        """List available templates"""
        templates = []
        
        for pleading_type, template_name in self.TEMPLATE_MAP.items():
            template_path = self.templates_dir / template_name
            templates.append({
                "type": pleading_type.value,
                "template": template_name,
                "available": template_path.exists(),
                "description": self._get_template_description(pleading_type)
            })
        
        return templates
    
    def _get_template_description(self, pleading_type: PleadingType) -> str:
        """Get description for pleading type"""
        descriptions = {
            PleadingType.VERIFIED_COMPLAINT: "Initial complaint filing with verification",
            PleadingType.MOTION_INTERIM_RELIEF: "Motion for interim/emergency relief",
            PleadingType.CERTIFICATION_SUPPORT: "Certification in support of motion",
        }
        return descriptions.get(pleading_type, "")


class FilingsSearchEngine:
    """
    Full-text search engine for filed documents.
    
    Uses SQLite FTS5 for fast keyword search across
    extracted document text.
    """
    
    def __init__(self, db_path: Optional[Path] = None):
        """
        Initialize search engine.
        
        Args:
            db_path: Path to SQLite database
        """
        self.db_path = db_path
        self._connection = None
    
    def extract_text_from_pdf(self, pdf_path: Path) -> str:
        """Extract text from PDF"""
        if PdfReader is None:
            raise RuntimeError("pypdf not available")
        
        reader = PdfReader(str(pdf_path))
        parts = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                parts.append("")
        
        text = "\n".join(parts)
        return self._normalize_text(text)
    
    def extract_text_from_docx(self, docx_path: Path) -> str:
        """Extract text from DOCX"""
        if Document is None:
            raise RuntimeError("python-docx not available")
        
        doc = Document(str(docx_path))
        parts = []
        
        for p in doc.paragraphs:
            if p.text:
                parts.append(p.text)
        
        # Tables
        for t in doc.tables:
            for row in t.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)
        
        text = "\n".join(parts)
        return self._normalize_text(text)
    
    def _normalize_text(self, text: str) -> str:
        """Normalize extracted text"""
        text = re.sub(r"Page\s+\d+\s+of\s+\d+", "", text, flags=re.I)
        text = re.sub(r"\s+\n", "\n", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()
    
    def build_index(self, txt_dir: Path, db_path: Path) -> int:
        """
        Build FTS5 index from text files.
        
        Args:
            txt_dir: Directory containing .txt files
            db_path: Output database path
            
        Returns:
            Number of documents indexed
        """
        if db_path.exists():
            db_path.unlink()
        
        con = sqlite3.connect(str(db_path))
        cur = con.cursor()
        cur.execute("PRAGMA journal_mode=WAL;")
        cur.execute("CREATE VIRTUAL TABLE docs USING fts5(path, title, body);")
        
        count = 0
        for p in sorted(txt_dir.rglob("*.txt")):
            if p.name.endswith(".jsonl"):
                continue
            
            body = p.read_text(encoding="utf-8", errors="ignore")
            cur.execute(
                "INSERT INTO docs(path, title, body) VALUES(?, ?, ?);",
                (str(p), p.name, body)
            )
            count += 1
        
        con.commit()
        con.close()
        
        self.db_path = db_path
        return count
    
    def search(
        self,
        query: str,
        limit: int = 20
    ) -> List[SearchResult]:
        """
        Search the FTS index.
        
        Args:
            query: FTS5 query (supports NEAR, AND, OR, etc.)
            limit: Maximum results
            
        Returns:
            List of search results
        """
        if self.db_path is None or not self.db_path.exists():
            return []
        
        con = sqlite3.connect(str(self.db_path))
        cur = con.cursor()
        
        try:
            cur.execute(
                """
                SELECT path, title, 
                       snippet(docs, 2, '[', ']', ' … ', 12) as snip,
                       rank
                FROM docs 
                WHERE docs MATCH ? 
                ORDER BY rank
                LIMIT ?;
                """,
                (query, limit)
            )
            rows = cur.fetchall()
        except sqlite3.OperationalError:
            rows = []
        finally:
            con.close()
        
        return [
            SearchResult(
                document=row[0],
                snippet=row[2],
                score=abs(row[3]) if row[3] else 0,
                context=row[1]
            )
            for row in rows
        ]
    
    def grep_search(
        self,
        txt_dir: Path,
        pattern: str,
        literal: bool = False,
        context: int = 80,
        max_hits: int = 50
    ) -> List[SearchResult]:
        """
        Regex search without database.
        
        Args:
            txt_dir: Directory to search
            pattern: Regex or literal pattern
            literal: If True, escape regex
            context: Context characters around match
            max_hits: Maximum matches
            
        Returns:
            List of search results
        """
        rx = re.compile(re.escape(pattern) if literal else pattern, flags=re.I)
        results = []
        
        for p in sorted(txt_dir.rglob("*.txt")):
            if p.name.endswith(".jsonl"):
                continue
            
            text = p.read_text(encoding="utf-8", errors="ignore")
            
            for m in rx.finditer(text):
                if len(results) >= max_hits:
                    return results
                
                s = max(0, m.start() - context)
                e = min(len(text), m.end() + context)
                snippet = text[s:e].replace("\n", " ")
                
                results.append(SearchResult(
                    document=str(p),
                    snippet=f"...{snippet}...",
                    score=1.0,
                    context=m.group(0)
                ))
        
        return results


class ExhibitMerger:
    """Merge PDFs for exhibit packages"""
    
    @staticmethod
    def merge_pdfs(output_path: Path, input_pdfs: List[Path]) -> bool:
        """
        Merge multiple PDFs into one.
        
        Args:
            output_path: Output PDF path
            input_pdfs: List of input PDF paths
            
        Returns:
            Success status
        """
        if PdfReader is None or PdfWriter is None:
            raise RuntimeError("pypdf not available")
        
        writer = PdfWriter()
        
        for pdf_path in input_pdfs:
            reader = PdfReader(str(pdf_path))
            for page in reader.pages:
                writer.add_page(page)
        
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        with output_path.open("wb") as f:
            writer.write(f)
        
        return True
    
    @staticmethod
    def add_exhibit_labels(
        pdf_path: Path,
        exhibit_label: str,
        output_path: Path
    ) -> bool:
        """
        Add exhibit label to PDF pages.
        
        Args:
            pdf_path: Input PDF
            exhibit_label: Label (e.g., "EXHIBIT A")
            output_path: Output PDF
            
        Returns:
            Success status
        """
        # This would add watermarks/stamps - simplified for now
        if PdfReader is None or PdfWriter is None:
            raise RuntimeError("pypdf not available")
        
        # Copy the PDF (label addition would require reportlab)
        import shutil
        shutil.copy(pdf_path, output_path)
        return True


# Convenience functions
def generate_complaint(
    facts: Dict[str, Any],
    output_path: Optional[str] = None
) -> PleadingResult:
    """Generate verified complaint"""
    generator = NJCivilPleadingsGenerator()
    return generator.generate_pleading(
        PleadingType.VERIFIED_COMPLAINT,
        facts,
        Path(output_path) if output_path else None
    )


def generate_motion(
    facts: Dict[str, Any],
    output_path: Optional[str] = None
) -> PleadingResult:
    """Generate motion for interim relief"""
    generator = NJCivilPleadingsGenerator()
    return generator.generate_pleading(
        PleadingType.MOTION_INTERIM_RELIEF,
        facts,
        Path(output_path) if output_path else None
    )


def generate_certification(
    facts: Dict[str, Any],
    output_path: Optional[str] = None
) -> PleadingResult:
    """Generate certification in support"""
    generator = NJCivilPleadingsGenerator()
    return generator.generate_pleading(
        PleadingType.CERTIFICATION_SUPPORT,
        facts,
        Path(output_path) if output_path else None
    )


def search_filings(
    query: str,
    db_path: str,
    limit: int = 20
) -> List[Dict[str, Any]]:
    """Search indexed filings"""
    engine = FilingsSearchEngine(Path(db_path))
    results = engine.search(query, limit)
    return [
        {
            "document": r.document,
            "snippet": r.snippet,
            "score": r.score
        }
        for r in results
    ]
